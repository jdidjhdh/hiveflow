"""HiveFlow - RAG (Retrieval-Augmented Generation) Module

Provides:
- DocumentProcessor: Parse PDF/Text/Markdown/HTML/DOCX documents
- TextChunker: Multiple chunking strategies (fixed, recursive, semantic, markdown-aware)
- VectorStore: Unified interface for Chroma/FAISS/Memory vector storage
- RAGPipeline: Complete retrieval pipeline with reranking and context assembly
- KnowledgeBase: Full knowledge base management with metadata filtering

Integrates with HiveFlow's blackboard system and LLM client.
"""
import hashlib
import json
import logging
import os
import re
import time
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from enum import Enum
from typing import Any, Callable, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


# ======================== Document Processor ========================

class DocumentType(str, Enum):
    """Supported document types."""
    TEXT = "text"
    MARKDOWN = "markdown"
    HTML = "html"
    PDF = "pdf"
    DOCX = "docx"
    CSV = "csv"
    JSON = "json"


@dataclass
class Document:
    """A parsed document with metadata."""
    doc_id: str
    content: str
    doc_type: DocumentType
    metadata: Dict[str, Any] = field(default_factory=dict)
    chunks: List["DocumentChunk"] = field(default_factory=list)

    @staticmethod
    def compute_doc_id(content: str, source: str = "") -> str:
        """Compute a stable document ID from content."""
        raw = f"{source}:{content}" if source else content
        return hashlib.md5(raw.encode()).hexdigest()[:12]


class DocumentProcessor:
    """
    Parse documents from various formats into standardized Document objects.
    
    Usage:
        processor = DocumentProcessor()
        doc = processor.parse("file.pdf")
        doc = processor.parse_text("Hello world", source="manual")
        doc = processor.parse_markdown("# Title\nContent")
    """

    def parse(self, file_path: str, doc_type: Optional[DocumentType] = None, **kwargs) -> Document:
        """Parse a file into a Document."""
        if doc_type is None:
            doc_type = self._detect_type(file_path)

        with open(file_path, "rb") as f:
            raw = f.read()

        if doc_type == DocumentType.PDF:
            return self._parse_pdf(raw, file_path, **kwargs)
        elif doc_type == DocumentType.HTML:
            return self._parse_html(raw, file_path, **kwargs)
        elif doc_type == DocumentType.DOCX:
            return self._parse_docx(raw, file_path, **kwargs)
        elif doc_type == DocumentType.CSV:
            return self._parse_csv(raw, file_path, **kwargs)
        elif doc_type == DocumentType.JSON:
            return self._parse_json(raw, file_path, **kwargs)
        else:
            # Default: text/markdown
            content = raw.decode("utf-8", errors="replace")
            is_md = file_path.endswith((".md", ".markdown"))
            return Document(
                doc_id=Document.compute_doc_id(content, file_path),
                content=content,
                doc_type=DocumentType.MARKDOWN if is_md else DocumentType.TEXT,
                metadata={"source": file_path, "size": len(content)},
            )

    def parse_text(self, text: str, source: str = "") -> Document:
        """Parse raw text."""
        return Document(
            doc_id=Document.compute_doc_id(text, source),
            content=text,
            doc_type=DocumentType.TEXT,
            metadata={"source": source, "size": len(text)},
        )

    def parse_markdown(self, markdown: str, source: str = "") -> Document:
        """Parse markdown content."""
        return Document(
            doc_id=Document.compute_doc_id(markdown, source),
            content=markdown,
            doc_type=DocumentType.MARKDOWN,
            metadata={"source": source, "size": len(markdown)},
        )

    def parse_html(self, html: str, source: str = "") -> Document:
        """Parse HTML content."""
        text = self._strip_html(html)
        return Document(
            doc_id=Document.compute_doc_id(text, source),
            content=text,
            doc_type=DocumentType.HTML,
            metadata={"source": source, "size": len(text), "original_html_size": len(html)},
        )

    def _detect_type(self, file_path: str) -> DocumentType:
        ext = os.path.splitext(file_path)[1].lower()
        mapping = {
            ".txt": DocumentType.TEXT,
            ".md": DocumentType.MARKDOWN,
            ".markdown": DocumentType.MARKDOWN,
            ".html": DocumentType.HTML,
            ".htm": DocumentType.HTML,
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".csv": DocumentType.CSV,
            ".json": DocumentType.JSON,
        }
        return mapping.get(ext, DocumentType.TEXT)

    def _strip_html(self, html: str) -> str:
        """Strip HTML tags, keeping text content."""
        # Remove script/style
        html = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
        html = re.sub(r"<style[^>]*>.*?</style>", "", html, flags=re.DOTALL | re.IGNORECASE)
        # Remove all tags
        text = re.sub(r"<[^>]+>", " ", html)
        # Normalize whitespace
        text = re.sub(r"\s+", " ", text).strip()
        return text

    def _parse_pdf(self, raw: bytes, path: str, **kwargs) -> Document:
        """Parse PDF file."""
        try:
            import pypdf
            from io import BytesIO
            reader = pypdf.PdfReader(BytesIO(raw))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            try:
                import pdfplumber
                from io import BytesIO
                with pdfplumber.open(BytesIO(raw)) as pdf:
                    text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            except ImportError:
                # Fallback: extract ASCII text only
                text = raw.decode("utf-8", errors="ignore")
                logger.warning("No PDF library available, using raw text extraction")

        return Document(
            doc_id=Document.compute_doc_id(text, path),
            content=text,
            doc_type=DocumentType.PDF,
            metadata={"source": path, "size": len(text)},
        )

    def _parse_html(self, raw: bytes, path: str, **kwargs) -> Document:
        """Parse HTML file."""
        html = raw.decode("utf-8", errors="replace")
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "html.parser")
            text = soup.get_text(separator="\n", strip=True)
        except ImportError:
            text = self._strip_html(html)

        return Document(
            doc_id=Document.compute_doc_id(text, path),
            content=text,
            doc_type=DocumentType.HTML,
            metadata={"source": path, "size": len(text)},
        )

    def _parse_docx(self, raw: bytes, path: str, **kwargs) -> Document:
        """Parse DOCX file."""
        try:
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(raw))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            text = "[docx parsing requires python-docx]"
            logger.warning("python-docx not available for DOCX parsing")

        return Document(
            doc_id=Document.compute_doc_id(text, path),
            content=text,
            doc_type=DocumentType.DOCX,
            metadata={"source": path, "size": len(text)},
        )

    def _parse_csv(self, raw: bytes, path: str, **kwargs) -> Document:
        """Parse CSV file into text format."""
        import csv
        from io import StringIO
        text = raw.decode("utf-8", errors="replace")
        reader = csv.reader(StringIO(text))
        rows = list(reader)
        # Convert to readable text: header + rows
        if not rows:
            content = ""
        else:
            headers = rows[0]
            content = "\n".join(
                ", ".join(f"{h}: {row[i] if i < len(row) else ''}" for i, h in enumerate(headers))
                for row in rows[1:]
            )

        return Document(
            doc_id=Document.compute_doc_id(content, path),
            content=content,
            doc_type=DocumentType.CSV,
            metadata={"source": path, "rows": len(rows) - 1, "columns": len(rows[0]) if rows else 0},
        )

    def _parse_json(self, raw: bytes, path: str, **kwargs) -> Document:
        """Parse JSON file."""
        text = raw.decode("utf-8", errors="replace")
        try:
            data = json.loads(text)
            # Flatten JSON to readable text
            content = self._flatten_json(data)
        except json.JSONDecodeError:
            content = text

        return Document(
            doc_id=Document.compute_doc_id(content, path),
            content=content,
            doc_type=DocumentType.JSON,
            metadata={"source": path, "size": len(content)},
        )

    @staticmethod
    def _flatten_json(obj: Any, prefix: str = "") -> str:
        """Flatten JSON to readable text."""
        lines = []
        if isinstance(obj, dict):
            for k, v in obj.items():
                key = f"{prefix}.{k}" if prefix else k
                if isinstance(v, (dict, list)):
                    lines.append(DocumentProcessor._flatten_json(v, key))
                else:
                    lines.append(f"{key}: {v}")
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                key = f"{prefix}[{i}]"
                if isinstance(item, (dict, list)):
                    lines.append(DocumentProcessor._flatten_json(item, key))
                else:
                    lines.append(f"{key}: {item}")
        else:
            lines.append(f"{prefix}: {obj}")
        return "\n".join(lines)


# ======================== Text Chunker ========================

@dataclass
class DocumentChunk:
    """A chunk of text from a document."""
    chunk_id: str
    doc_id: str
    content: str
    index: int
    metadata: Dict[str, Any] = field(default_factory=dict)

    def compute_chunk_id(self) -> str:
        """Compute stable chunk ID."""
        return hashlib.md5(f"{self.doc_id}:{self.index}:{self.content}".encode()).hexdigest()[:12]


class ChunkStrategy(str, Enum):
    """Available chunking strategies."""
    FIXED = "fixed"              # Fixed size chunks
    RECURSIVE = "recursive"      # Recursive character splitting
    SEMANTIC = "semantic"        # Paragraph-based splitting
    MARKDOWN = "markdown"        # Markdown-aware splitting
    CODE = "code"                # Code-aware splitting


class TextChunker:
    """
    Split documents into chunks for vector embedding.
    
    Usage:
        chunker = TextChunker(strategy="recursive", chunk_size=500, chunk_overlap=50)
        chunks = chunker.chunk(document)
    """

    def __init__(
        self,
        strategy: ChunkStrategy = ChunkStrategy.RECURSIVE,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        separators: Optional[List[str]] = None,
    ):
        self.strategy = strategy
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or self._default_separators()

    def chunk(self, doc: Document) -> List[DocumentChunk]:
        """Split a document into chunks."""
        if self.strategy == ChunkStrategy.FIXED:
            return self._fixed_chunk(doc)
        elif self.strategy == ChunkStrategy.RECURSIVE:
            return self._recursive_chunk(doc)
        elif self.strategy == ChunkStrategy.SEMANTIC:
            return self._semantic_chunk(doc)
        elif self.strategy == ChunkStrategy.MARKDOWN:
            return self._markdown_chunk(doc)
        elif self.strategy == ChunkStrategy.CODE:
            return self._code_chunk(doc)
        else:
            return self._recursive_chunk(doc)

    def _fixed_chunk(self, doc: Document) -> List[DocumentChunk]:
        """Split by fixed character count."""
        text = doc.content
        chunks = []
        start = 0
        idx = 0
        while start < len(text):
            end = start + self.chunk_size
            chunk_text = text[start:end]
            chunk = DocumentChunk(
                chunk_id="",
                doc_id=doc.doc_id,
                content=chunk_text,
                index=idx,
                metadata={**doc.metadata, "chunk_strategy": "fixed"},
            )
            chunk.chunk_id = chunk.compute_chunk_id()
            chunks.append(chunk)
            start = end - self.chunk_overlap
            idx += 1
        return chunks

    def _recursive_chunk(self, doc: Document) -> List[DocumentChunk]:
        """Recursive character splitting with multiple separators."""
        text = doc.content
        chunks = []
        self._recursive_split(text, self.separators, chunks, doc, 0)
        return chunks

    def _recursive_split(self, text: str, separators: List[str], chunks: List[DocumentChunk], doc: Document, base_idx: int) -> int:
        """Recursively split text."""
        if len(text) <= self.chunk_size:
            if text.strip():
                chunk = DocumentChunk(
                    chunk_id="", doc_id=doc.doc_id, content=text,
                    index=base_idx, metadata={**doc.metadata, "chunk_strategy": "recursive"},
                )
                chunk.chunk_id = chunk.compute_chunk_id()
                chunks.append(chunk)
            return base_idx + 1

        separator = ""
        for sep in separators:
            if sep in text:
                separator = sep
                break

        if not separator:
            # Fallback to fixed chunk
            return self._fixed_chunk_text(text, chunks, doc, base_idx)

        parts = text.split(separator)
        current = ""
        idx = base_idx
        for part in parts:
            if len(current) + len(part) + len(separator) <= self.chunk_size:
                current += (separator if current else "") + part
            else:
                if current.strip():
                    chunk = DocumentChunk(
                        chunk_id="", doc_id=doc.doc_id, content=current,
                        index=idx, metadata={**doc.metadata, "chunk_strategy": "recursive"},
                    )
                    chunk.chunk_id = chunk.compute_chunk_id()
                    chunks.append(chunk)
                    idx += 1
                # Start with overlap
                if self.chunk_overlap > 0 and len(current) > self.chunk_overlap:
                    current = current[-self.chunk_overlap:] + (separator if current.strip() else "") + part
                else:
                    current = part

        if current.strip():
            chunk = DocumentChunk(
                chunk_id="", doc_id=doc.doc_id, content=current,
                index=idx, metadata={**doc.metadata, "chunk_strategy": "recursive"},
            )
            chunk.chunk_id = chunk.compute_chunk_id()
            chunks.append(chunk)
            idx += 1

        return idx

    def _fixed_chunk_text(self, text: str, chunks: List[DocumentChunk], doc: Document, idx: int) -> int:
        """Fallback fixed chunking."""
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            chunk_text = text[start:end]
            chunk = DocumentChunk(
                chunk_id="", doc_id=doc.doc_id, content=chunk_text,
                index=idx, metadata={**doc.metadata, "chunk_strategy": "recursive"},
            )
            chunk.chunk_id = chunk.compute_chunk_id()
            chunks.append(chunk)
            start = end - self.chunk_overlap
            idx += 1
        return idx

    def _semantic_chunk(self, doc: Document) -> List[DocumentChunk]:
        """Split by paragraphs, merge small ones."""
        paragraphs = re.split(r"\n\n+", doc.content)
        chunks = []
        current = ""
        idx = 0
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            if len(current) + len(para) <= self.chunk_size and current:
                current += "\n\n" + para
            else:
                if current.strip():
                    chunk = DocumentChunk(
                        chunk_id="", doc_id=doc.doc_id, content=current,
                        index=idx, metadata={**doc.metadata, "chunk_strategy": "semantic"},
                    )
                    chunk.chunk_id = chunk.compute_chunk_id()
                    chunks.append(chunk)
                    idx += 1
                current = para

        if current.strip():
            chunk = DocumentChunk(
                chunk_id="", doc_id=doc.doc_id, content=current,
                index=idx, metadata={**doc.metadata, "chunk_strategy": "semantic"},
            )
            chunk.chunk_id = chunk.compute_chunk_id()
            chunks.append(chunk)

        return chunks

    def _markdown_chunk(self, doc: Document) -> List[DocumentChunk]:
        """Split by markdown headings."""
        lines = doc.content.split("\n")
        chunks = []
        current_lines = []
        idx = 0

        for line in lines:
            if re.match(r"^#{1,6}\s+", line) and current_lines:
                # New heading starts a new chunk
                chunk_text = "\n".join(current_lines)
                if chunk_text.strip():
                    chunk = DocumentChunk(
                        chunk_id="", doc_id=doc.doc_id, content=chunk_text,
                        index=idx, metadata={**doc.metadata, "chunk_strategy": "markdown"},
                    )
                    chunk.chunk_id = chunk.compute_chunk_id()
                    chunks.append(chunk)
                    idx += 1
                current_lines = [line]
            else:
                current_lines.append(line)

        if current_lines:
            chunk_text = "\n".join(current_lines)
            if chunk_text.strip():
                chunk = DocumentChunk(
                    chunk_id="", doc_id=doc.doc_id, content=chunk_text,
                    index=idx, metadata={**doc.metadata, "chunk_strategy": "markdown"},
                )
                chunk.chunk_id = chunk.compute_chunk_id()
                chunks.append(chunk)

        return chunks

    def _code_chunk(self, doc: Document) -> List[DocumentChunk]:
        """Split code by class/function definitions."""
        # Split by class and function definitions
        pattern = r"(^(class|def)\s+\w+)"
        parts = re.split(pattern, doc.content, flags=re.MULTILINE)

        chunks = []
        current = ""
        idx = 0

        for i, part in enumerate(parts):
            if re.match(r"^(class|def)\s+\w+", part) and current:
                if current.strip():
                    chunk = DocumentChunk(
                        chunk_id="", doc_id=doc.doc_id, content=current,
                        index=idx, metadata={**doc.metadata, "chunk_strategy": "code"},
                    )
                    chunk.chunk_id = chunk.compute_chunk_id()
                    chunks.append(chunk)
                    idx += 1
                current = part
            else:
                current += part

        if current.strip():
            chunk = DocumentChunk(
                chunk_id="", doc_id=doc.doc_id, content=current,
                index=idx, metadata={**doc.metadata, "chunk_strategy": "code"},
            )
            chunk.chunk_id = chunk.compute_chunk_id()
            chunks.append(chunk)

        return chunks

    def _default_separators(self) -> List[str]:
        """Default separators for recursive chunking."""
        return ["\n\n\n", "\n\n", "\n", ". ", ", ", " ", ""]


# ======================== Vector Store ========================

@dataclass
class SearchResult:
    """A single search result from a vector store."""
    chunk: DocumentChunk
    score: float
    metadata: Dict[str, Any] = field(default_factory=dict)


class VectorStore(ABC):
    """Abstract base class for vector storage."""

    @abstractmethod
    async def add(self, chunks: List[DocumentChunk], embeddings: List[List[float]]) -> List[str]:
        """Add chunks with embeddings. Returns chunk IDs."""
        ...

    @abstractmethod
    async def search(
        self,
        query_embedding: List[float],
        top_k: int = 5,
        filters: Optional[Dict[str, Any]] = None,
    ) -> List[SearchResult]:
        """Search for similar chunks."""
        ...

    @abstractmethod
    async def delete(self, doc_id: str) -> int:
        """Delete all chunks for a document. Returns count deleted."""
        ...

    @abstractmethod
    async def count(self) -> int:
        """Total number of chunks stored."""
        ...


class MemoryVectorStore(VectorStore):
    """In-memory vector store using cosine similarity. Good for testing/small datasets."""

    def __init__(self):
        self._chunks: Dict[str, Tuple[DocumentChunk, List[float]]] = {}

    async def add(self, chunks: List[DocumentChunk], embeddings: List[List[float]]) -> List[str]:
        ids = []
        for chunk, embedding in zip(chunks, embeddings):
            self._chunks[chunk.chunk_id] = (chunk, embedding)
            ids.append(chunk.chunk_id)
        return ids

    async def search(
        self,
        query_embedding: List[float],
        top_k: int = 5,
        filters: Optional[Dict[str, Any]] = None,
    ) -> List[SearchResult]:
        results = []
        for chunk_id, (chunk, embedding) in self._chunks.items():
            # Apply filters
            if filters and not self._matches_filters(chunk.metadata, filters):
                continue

            score = self._cosine_similarity(query_embedding, embedding)
            results.append(SearchResult(chunk=chunk, score=score))

        results.sort(key=lambda r: r.score, reverse=True)
        return results[:top_k]

    async def delete(self, doc_id: str) -> int:
        to_delete = [cid for cid, (chunk, _) in self._chunks.items() if chunk.doc_id == doc_id]
        for cid in to_delete:
            del self._chunks[cid]
        return len(to_delete)

    async def count(self) -> int:
        return len(self._chunks)

    @staticmethod
    def _cosine_similarity(a: List[float], b: List[float]) -> float:
        """Compute cosine similarity between two vectors."""
        if not a or not b:
            return 0.0
        dot = sum(x * y for x, y in zip(a, b))
        norm_a = sum(x * x for x in a) ** 0.5
        norm_b = sum(x * x for x in b) ** 0.5
        if norm_a == 0 or norm_b == 0:
            return 0.0
        return dot / (norm_a * norm_b)

    @staticmethod
    def _matches_filters(metadata: Dict[str, Any], filters: Dict[str, Any]) -> bool:
        """Check if metadata matches all filters."""
        for key, value in filters.items():
            if key not in metadata:
                return False
            if isinstance(value, (list, tuple)):
                if metadata[key] not in value:
                    return False
            elif metadata[key] != value:
                return False
        return True


try:
    import chromadb
    _CHROMA_AVAILABLE = True
except ImportError:
    _CHROMA_AVAILABLE = False


class ChromaVectorStore(VectorStore):
    """Chroma-backed vector store. Requires chromadb package."""

    def __init__(
        self,
        collection_name: str = "hiveflow_kb",
        persist_dir: Optional[str] = None,
        embedding_fn: Optional[Callable[[List[str]], List[List[float]]]] = None,
    ):
        if not _CHROMA_AVAILABLE:
            raise ImportError("chromadb is required for ChromaVectorStore")

        if persist_dir:
            client = chromadb.PersistentClient(path=persist_dir)
        else:
            client = chromadb.EphemeralClient()

        self.collection = client.get_or_create_collection(
            name=collection_name,
            metadata={"hnsw:space": "cosine"},
        )
        self.embedding_fn = embedding_fn

    async def add(self, chunks: List[DocumentChunk], embeddings: List[List[float]]) -> List[str]:
        if self.embedding_fn and not embeddings:
            texts = [c.content for c in chunks]
            embeddings = self.embedding_fn(texts)

        documents = [c.content for c in chunks]
        ids = [c.chunk_id for c in chunks]
        metadatas = []
        for c in chunks:
            meta = {k: str(v) if not isinstance(v, str) else v for k, v in c.metadata.items()}
            meta["doc_id"] = c.doc_id
            meta["index"] = c.index
            metadatas.append(meta)

        self.collection.add(
            ids=ids,
            documents=documents,
            embeddings=embeddings if embeddings else None,
            metadatas=metadatas,
        )
        return ids

    async def search(
        self,
        query_embedding: List[float],
        top_k: int = 5,
        filters: Optional[Dict[str, Any]] = None,
    ) -> List[SearchResult]:
        where = None
        if filters:
            where = {
                "$and": [
                    {k: {"$eq": v} if not isinstance(v, (list, tuple)) else {k: {"$in": list(v)}}}
                    for k, v in filters.items()
                ]
            }

        results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=top_k,
            where=where,
            include=["documents", "metadatas", "distances"],
        )

        search_results = []
        if results["ids"] and results["ids"][0]:
            for i, (chunk_id, doc, meta, distance) in enumerate(zip(
                results["ids"][0],
                results["documents"][0],
                results["metadatas"][0],
                results["distances"][0],
            )):
                chunk = DocumentChunk(
                    chunk_id=chunk_id,
                    doc_id=meta.get("doc_id", ""),
                    content=doc or "",
                    index=meta.get("index", 0),
                    metadata=meta or {},
                )
                # Convert distance to similarity
                score = 1.0 - distance
                search_results.append(SearchResult(chunk=chunk, score=score))

        return search_results

    async def delete(self, doc_id: str) -> int:
        existing = self.collection.get(where={"doc_id": {"$eq": doc_id}})
        if not existing["ids"]:
            return 0
        count = len(existing["ids"])
        self.collection.delete(ids=existing["ids"])
        return count

    async def count(self) -> int:
        return self.collection.count()


# ======================== RAG Pipeline ========================

@dataclass
class RAGResult:
    """Result of a RAG query."""
    query: str
    answer: str
    sources: List[SearchResult] = field(default_factory=list)
    context: str = ""
    latency_ms: float = 0.0


class EmbeddingModel(ABC):
    """Abstract embedding model."""

    @abstractmethod
    async def embed(self, texts: List[str]) -> List[List[float]]:
        ...

    @abstractmethod
    async def embed_query(self, text: str) -> List[float]:
        ...


class DummyEmbeddingModel(EmbeddingModel):
    """Dummy embedding that returns random-ish vectors. For testing only."""

    def __init__(self, dim: int = 128):
        self.dim = dim

    def _hash_vector(self, text: str) -> List[float]:
        """Generate a deterministic vector from text hash."""
        import random
        seed = hash(text) & 0xFFFFFFFF
        rng = random.Random(seed)
        return [rng.gauss(0, 1) for _ in range(self.dim)]

    async def embed(self, texts: List[str]) -> List[List[float]]:
        return [self._hash_vector(t) for t in texts]

    async def embed_query(self, text: str) -> List[float]:
        return self._hash_vector(text)


class RAGPipeline:
    """
    Complete RAG pipeline: retrieve -> rerank -> generate answer.
    
    Usage:
        pipeline = RAGPipeline(
            vector_store=MemoryVectorStore(),
            embedding_model=DummyEmbeddingModel(),
            llm_client=openai_client,
        )
        
        # Add documents
        await pipeline.add_document(doc, chunker=TextChunker())
        
        # Query
        result = await pipeline.query("What is the main topic?")
    """

    def __init__(
        self,
        vector_store: VectorStore,
        embedding_model: EmbeddingModel,
        llm_client=None,  # Optional: for answer generation
        model: str = "",
        top_k: int = 5,
        rerank_fn: Optional[Callable[[str, str], float]] = None,
    ):
        self.vector_store = vector_store
        self.embedding_model = embedding_model
        self.llm_client = llm_client
        self.model = model
        self.top_k = top_k
        self.rerank_fn = rerank_fn

    async def add_document(self, doc: Document, chunker: Optional[TextChunker] = None) -> List[str]:
        """Process and add a document to the knowledge base."""
        if chunker is None:
            chunker = TextChunker()

        chunks = chunker.chunk(doc)
        if not chunks:
            return []

        # Generate embeddings
        texts = [c.content for c in chunks]
        embeddings = await self.embedding_model.embed(texts)

        # Store in vector store
        ids = await self.vector_store.add(chunks, embeddings)
        doc.chunks = chunks
        logger.info(f"Added {len(ids)} chunks from document {doc.doc_id}")
        return ids

    async def query(
        self,
        query: str,
        filters: Optional[Dict[str, Any]] = None,
        top_k: int = 0,
        include_context: bool = True,
    ) -> RAGResult:
        """Execute a RAG query."""
        import time
        start = time.monotonic()
        k = top_k or self.top_k

        # Retrieve
        query_embedding = await self.embedding_model.embed_query(query)
        results = await self.vector_store.search(query_embedding, top_k=k * 2, filters=filters)

        # Rerank if function provided
        if self.rerank_fn and results:
            results = self._rerank(query, results)

        results = results[:k]

        # Build context
        context = "\n\n".join(
            f"[Source {i+1}] (score: {r.score:.3f})\n{r.chunk.content}"
            for i, r in enumerate(results)
        ) if results else "No relevant information found."

        # Generate answer
        answer = ""
        if self.llm_client:
            try:
                from . import LLMMessage
                messages = [
                    LLMMessage(role="system", content=(
                        "You are a helpful assistant. Answer the question based on the provided context. "
                        "If the context doesn't contain relevant information, say so."
                    )),
                    LLMMessage(role="user", content=f"Context:\n{context}\n\nQuestion: {query}"),
                ]
                response = await self.llm_client.chat(
                    messages=messages, model=self.model, temperature=0.1, max_tokens=1024,
                )
                answer = response.content
            except Exception as e:
                logger.warning(f"LLM answer generation failed: {e}")
                answer = f"[Error generating answer: {e}]"

        elapsed_ms = (time.monotonic() - start) * 1000

        return RAGResult(
            query=query,
            answer=answer,
            sources=results,
            context=context if include_context else "",
            latency_ms=elapsed_ms,
        )

    async def search(
        self,
        query: str,
        filters: Optional[Dict[str, Any]] = None,
        top_k: int = 0,
    ) -> List[SearchResult]:
        """Search without generating an answer (pure retrieval)."""
        k = top_k or self.top_k
        query_embedding = await self.embedding_model.embed_query(query)
        results = await self.vector_store.search(query_embedding, top_k=k, filters=filters)

        if self.rerank_fn and results:
            results = self._rerank(query, results)

        return results[:k]

    def _rerank(self, query: str, results: List[SearchResult]) -> List[SearchResult]:
        """Rerank results using provided rerank function."""
        reranked = []
        for r in results:
            score = self.rerank_fn(query, r.chunk.content)
            reranked.append(SearchResult(
                chunk=r.chunk,
                score=score,
                metadata=r.metadata,
            ))
        reranked.sort(key=lambda r: r.score, reverse=True)
        return reranked

    async def get_stats(self) -> Dict[str, Any]:
        """Get knowledge base statistics."""
        return {
            "total_chunks": await self.vector_store.count(),
        }


# ======================== KnowledgeBase Manager ========================

@dataclass
class KnowledgeBase:
    """A knowledge base with its documents."""
    kb_id: str
    name: str
    description: str = ""
    doc_count: int = 0
    chunk_count: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)
    created_at: float = field(default_factory=time.time)
    updated_at: float = field(default_factory=time.time)


class KnowledgeBaseManager:
    """
    Manages multiple knowledge bases.
    
    Usage:
        mgr = KnowledgeBaseManager()
        kb = await mgr.create_kb("my_kb", "My Knowledge Base")
        await mgr.add_document(kb.kb_id, doc)
        results = await mgr.query(kb.kb_id, "What is X?")
    """

    def __init__(
        self,
        vector_store: Optional[VectorStore] = None,
        embedding_model: Optional[EmbeddingModel] = None,
    ):
        self._knowledge_bases: Dict[str, KnowledgeBase] = {}
        self._pipelines: Dict[str, RAGPipeline] = {}

        # Use defaults if not provided
        if vector_store is None:
            vector_store = MemoryVectorStore()
        if embedding_model is None:
            embedding_model = DummyEmbeddingModel()

        self._vector_store = vector_store
        self._embedding_model = embedding_model

    async def create_kb(
        self,
        kb_id: str,
        name: str,
        description: str = "",
    ) -> KnowledgeBase:
        """Create a new knowledge base."""
        kb = KnowledgeBase(kb_id=kb_id, name=name, description=description)
        self._knowledge_bases[kb_id] = kb

        pipeline = RAGPipeline(
            vector_store=self._vector_store,
            embedding_model=self._embedding_model,
        )
        self._pipelines[kb_id] = pipeline

        return kb

    async def delete_kb(self, kb_id: str) -> bool:
        """Delete a knowledge base."""
        if kb_id in self._knowledge_bases:
            del self._knowledge_bases[kb_id]
            self._pipelines.pop(kb_id, None)
            return True
        return False

    async def list_kbs(self) -> List[KnowledgeBase]:
        """List all knowledge bases."""
        return list(self._knowledge_bases.values())

    async def add_document(self, kb_id: str, doc: Document, chunker: Optional[TextChunker] = None) -> List[str]:
        """Add a document to a knowledge base."""
        pipeline = self._pipelines.get(kb_id)
        if not pipeline:
            raise ValueError(f"Knowledge base not found: {kb_id}")

        ids = await pipeline.add_document(doc, chunker)
        kb = self._knowledge_bases[kb_id]
        kb.doc_count += 1
        kb.chunk_count += len(ids)
        kb.updated_at = time.time()
        return ids

    async def remove_document(self, kb_id: str, doc_id: str) -> int:
        """Remove a document from a knowledge base."""
        pipeline = self._pipelines.get(kb_id)
        if not pipeline:
            raise ValueError(f"Knowledge base not found: {kb_id}")

        count = await self._vector_store.delete(doc_id)
        kb = self._knowledge_bases[kb_id]
        kb.doc_count = max(0, kb.doc_count - 1)
        kb.chunk_count = max(0, kb.chunk_count - count)
        kb.updated_at = time.time()
        return count

    async def query(
        self,
        kb_id: str,
        query: str,
        top_k: int = 5,
        filters: Optional[Dict[str, Any]] = None,
    ) -> RAGResult:
        """Query a knowledge base."""
        pipeline = self._pipelines.get(kb_id)
        if not pipeline:
            raise ValueError(f"Knowledge base not found: {kb_id}")

        return await pipeline.query(query, filters=filters, top_k=top_k)

    async def search(
        self,
        kb_id: str,
        query: str,
        top_k: int = 5,
        filters: Optional[Dict[str, Any]] = None,
    ) -> List[SearchResult]:
        """Search a knowledge base without answer generation."""
        pipeline = self._pipelines.get(kb_id)
        if not pipeline:
            raise ValueError(f"Knowledge base not found: {kb_id}")

        # No doc_id filter - search all chunks in the shared vector store
        # (In production, you'd use per-KB collections or prefix filtering)
        return await pipeline.search(query, filters=filters, top_k=top_k)
